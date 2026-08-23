#!/usr/bin/env python3
"""Build DOCX files for Agora Studio's Knowledge Base from the sibling
Whatsapp-Chatbot-Gov repo's scheme CSV -- read-only, never writes to that
repo.

Why DOCX, not the raw CSV: Studio's Knowledge Base only accepts PDF/DOCX
(20MB/file), confirmed against Agora's docs -- see docs/AGORA_SETUP.md and
docs/VOICE_AGENT_SYSTEM_PROMPT.md. This script filters the ~4,112 scheme
CSV down to the agreed scope (Agriculture, Education & Learning, and the
*narrow* Loan/Finance definition -- Banking / Financial Services and
Insurance only, not the broader Business & Entrepreneurship bucket) using
the `Categories (All)` column, then renders each scheme's already-clean
`Scheme_info_detail` markdown into a DOCX section, batched into files well
under the 20MB cap.

Usage:
    .venv/bin/python scripts/build_knowledge_base.py
    .venv/bin/python scripts/build_knowledge_base.py --batch-size 40
    .venv/bin/python scripts/build_knowledge_base.py --broad-finance   # include Business & Entrepreneurship too

Output: data/knowledge_base/kb_XX.docx (+ data/knowledge_base/manifest.json)
Upload every kb_*.docx file to Agora Studio's Actions tab > Knowledge Base.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from bolo.config import settings  # noqa: E402

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_H3_RE = re.compile(r"^###\s+(.*)")
_H2_RE = re.compile(r"^##\s+(.*)")
_H1_RE = re.compile(r"^#\s+(.*)")
_BULLET_RE = re.compile(r"^\s*[-*]\s+(.*)")
_NUM_RE = re.compile(r"^\s*\d+[.)]\s+(.*)")
_MD_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
# XML 1.0 disallows most control characters (source data has a few stray
# ones); keep tab/newline/CR, drop everything else in that range.
_CONTROL_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


def _write_runs(paragraph, text: str) -> None:
    """Split on **bold** markers and add runs accordingly. Also flattens
    markdown links to 'text (url)' since docx hyperlinks need extra XML
    plumbing we don't need for a searchable knowledge base."""
    text = _CONTROL_CHARS_RE.sub("", text)
    text = _MD_LINK_RE.sub(r"\1 (\2)", text)
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        run = paragraph.add_run(m.group(1))
        run.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_markdown(doc: Document, markdown_text: str) -> None:
    if not isinstance(markdown_text, str) or not markdown_text.strip():
        return
    for raw_line in markdown_text.split("\n"):
        line = raw_line.rstrip()
        if not line.strip() or line.strip() == "---":
            continue
        if m := _H3_RE.match(line):
            doc.add_heading(m.group(1).strip(), level=3)
            continue
        if m := _H2_RE.match(line):
            doc.add_heading(m.group(1).strip(), level=2)
            continue
        if m := _H1_RE.match(line):
            doc.add_heading(m.group(1).strip(), level=2)  # demote -- H1 reserved for scheme name
            continue
        if m := _BULLET_RE.match(line):
            p = doc.add_paragraph(style="List Bullet")
            _write_runs(p, m.group(1).strip())
            continue
        if m := _NUM_RE.match(line):
            p = doc.add_paragraph(style="List Number")
            _write_runs(p, m.group(1).strip())
            continue
        p = doc.add_paragraph()
        _write_runs(p, line.strip())


def add_scheme(doc: Document, row: pd.Series) -> None:
    doc.add_heading(str(row.get("Scheme Name", "Untitled Scheme")), level=1)

    meta_bits = []
    for label, col in [
        ("Ministry", "Nodal Ministry"),
        ("State", "Beneficiary State"),
        ("Categories", "Categories (All)"),
        ("Open", "Open Date"),
        ("Close", "Close Date"),
        ("Scheme For", "Scheme For"),
    ]:
        val = row.get(col)
        if isinstance(val, str) and val.strip():
            meta_bits.append(f"{label}: {val.strip()}")
    if meta_bits:
        p = doc.add_paragraph(" | ".join(meta_bits))
        for run in p.runs:
            run.italic = True
            run.font.size = Pt(9)

    detail = row.get("Scheme_info_detail")
    if isinstance(detail, str) and detail.strip():
        # Skip the leading "# Scheme Name" line in the source markdown --
        # we already added our own H1 above with cleaner metadata.
        lines = detail.split("\n")
        if lines and lines[0].strip().startswith("# "):
            lines = lines[1:]
        add_markdown(doc, "\n".join(lines))
    else:
        # Fall back to the individual columns if Scheme_info_detail is blank.
        for heading, col in [
            ("Details", "Details (markdown)"),
            ("Benefits", "Benefits (markdown)"),
            ("Eligibility", "Eligibility (markdown)"),
            ("Application Process", "Application Process (markdown)"),
            ("Documents Required", "Documents Required (markdown)"),
            ("FAQs", "FAQs (markdown)"),
        ]:
            val = row.get(col)
            if isinstance(val, str) and val.strip():
                doc.add_heading(heading, level=2)
                add_markdown(doc, val)

    scheme_url = row.get("Scheme URL")
    if isinstance(scheme_url, str) and scheme_url.strip():
        doc.add_paragraph(f"Source: {scheme_url.strip()}")

    doc.add_page_break()


def build(csv_path: Path, out_dir: Path, batch_size: int, broad_finance: bool) -> None:
    df = pd.read_csv(csv_path)
    cats = df["Categories (All)"].fillna("")

    agri = cats.str.contains("Agriculture", case=False)
    edu = cats.str.contains("Education & Learning", case=False)
    finance_pattern = (
        "Banking|Financial Services and Insurance|Business & Entrepreneurship"
        if broad_finance
        else "Banking|Financial Services and Insurance"
    )
    finance = cats.str.contains(finance_pattern, case=False, regex=True)

    scope = agri | edu | finance
    subset = df[scope].reset_index(drop=True)

    print(f"Total schemes in source CSV: {len(df)}")
    print(f"  Agriculture matches:       {agri.sum()}")
    print(f"  Education matches:         {edu.sum()}")
    print(f"  Loan/Finance matches:      {finance.sum()} ({'broad' if broad_finance else 'narrow'})")
    print(f"  Union (knowledge base):    {len(subset)}")

    out_dir.mkdir(parents=True, exist_ok=True)
    manifest = []
    n_batches = (len(subset) + batch_size - 1) // batch_size

    for batch_idx in range(n_batches):
        batch = subset.iloc[batch_idx * batch_size : (batch_idx + 1) * batch_size]
        doc = Document()
        for _, row in batch.iterrows():
            add_scheme(doc, row)
            manifest.append(
                {
                    "scheme_name": row.get("Scheme Name"),
                    "batch_file": f"kb_{batch_idx + 1:02d}.docx",
                }
            )
        out_path = out_dir / f"kb_{batch_idx + 1:02d}.docx"
        doc.save(str(out_path))
        size_mb = out_path.stat().st_size / (1024 * 1024)
        print(f"  wrote {out_path.name}: {len(batch)} schemes, {size_mb:.2f} MB")
        if size_mb > 18:
            print(f"    WARNING: close to Studio's 20MB/file limit -- lower --batch-size")

    (out_dir / "manifest.json").write_text(json.dumps(manifest, indent=2, ensure_ascii=False))
    print(f"\nDone. Upload every kb_*.docx in {out_dir} to Studio's Actions tab > Knowledge Base.")


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--batch-size", type=int, default=40, help="Schemes per DOCX file (default: 40)"
    )
    parser.add_argument(
        "--broad-finance",
        action="store_true",
        help="Also include Business & Entrepreneurship schemes (~2,679 total instead of ~2,050)",
    )
    parser.add_argument(
        "--out-dir",
        type=Path,
        default=Path(__file__).resolve().parent.parent / "data" / "knowledge_base",
    )
    args = parser.parse_args()

    csv_path = settings.govscheme_repo_path / "web-scraping" / "4000_data.csv"
    if not csv_path.is_file():
        raise SystemExit(
            f"Can't find {csv_path}. Check GOVSCHEME_REPO_PATH in .env points at your "
            "Whatsapp-Chatbot-Gov checkout."
        )

    build(csv_path, args.out_dir, args.batch_size, args.broad_finance)


if __name__ == "__main__":
    main()
